from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUTPUT_DIR / "manual_v2_assets"
OPTIMIZED_DIR = ASSET_DIR / "optimized"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
OUTPUT_PATH = OUTPUT_DIR / "电力能源系统图上建模平台使用说明书_V2.0_紧凑图文版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5B6573"
WHITE = "FFFFFF"
GOLD = "C58B2A"
RED = "9B1C1C"
GREEN = "2F6B4F"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


@dataclass(frozen=True)
class PresetTokens:
    page_size: str = "US Letter 8.5 x 11 in"
    margins: str = "1.0 in"
    header_footer_distance: str = "0.492 in"
    content_width_dxa: int = CONTENT_WIDTH_DXA
    body_font_latin: str = "Calibri"
    body_font_east_asia: str = "微软雅黑"
    body_size_pt: float = 11.0
    body_after_pt: float = 6.0
    body_line_spacing: float = 1.25
    h1_size_pt: float = 16.0
    h1_before_pt: float = 18.0
    h1_after_pt: float = 10.0
    h2_size_pt: float = 13.0
    h2_before_pt: float = 14.0
    h2_after_pt: float = 7.0
    h3_size_pt: float = 12.0
    h3_before_pt: float = 10.0
    h3_after_pt: float = 5.0
    list_marker_in: float = 0.187
    list_text_in: float = 0.375
    list_hanging_in: float = 0.188


TOKENS = PresetTokens()


def set_run_font(
    run,
    *,
    latin: str = TOKENS.body_font_latin,
    east_asia: str = TOKENS.body_font_east_asia,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_style_font(style, size: float, color: str | None = None, bold: bool | None = None):
    style.font.name = TOKENS.body_font_latin
    style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS.body_font_latin)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS.body_font_latin)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS.body_font_east_asia)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, TOKENS.body_size_pt, color="20242A")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS.body_after_pt)
    normal.paragraph_format.line_spacing = TOKENS.body_line_spacing
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (TOKENS.h1_size_pt, BLUE, TOKENS.h1_before_pt, TOKENS.h1_after_pt),
        "Heading 2": (TOKENS.h2_size_pt, BLUE, TOKENS.h2_before_pt, TOKENS.h2_after_pt),
        "Heading 3": (TOKENS.h3_size_pt, DARK_BLUE, TOKENS.h3_before_pt, TOKENS.h3_after_pt),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    set_style_font(caption, 9.0, color=MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("电力能源系统图上建模平台 | 使用说明书 V2.0"), size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("第 "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    set_run_font(fp.add_run(" 页 / 共 "), size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")
    set_run_font(fp.add_run(" 页"), size=8.5, color=MUTED)

    props = doc.core_properties
    props.title = "电力能源系统图上建模平台使用说明书"
    props.subject = "图上建模、类与元件定义、参数量测、SVG 导入导出和文件交换"
    props.author = "电力能源系统图上建模平台项目组"
    props.keywords = "图上建模, 类定义, 元件定义, SVG, E文件, 使用手册"
    props.comments = "紧凑图文版，适用于 main@1267a934"


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_idx, row in enumerate(table.rows):
        for col_idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=9.0)
        if row_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "C8D1DC", size: str = "6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold=False, color="20242A", align=WD_ALIGN_PARAGRAPH.LEFT, size=9.0):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    p.clear()
    set_run_font(p.add_run(str(text)), size=size, bold=bold, color=color)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int], caption: str | None = None):
    global TABLE_COUNTER
    if caption:
        TABLE_COUNTER += 1
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(4)
        cp.paragraph_format.keep_with_next = True
        set_run_font(cp.add_run(f"表 {TABLE_COUNTER}  {caption}"), size=9.0, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], str(value), align=WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER)
            if len(table.rows) % 2 == 1:
                shade_cell(row.cells[idx], "F9FAFC")
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id: int, num_format: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), num_format)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        ppr.append(ind)
        lvl.append(ppr)
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(900, "bullet", "•", "Symbol")
    add_abstract(901, "decimal", "%1.")
    return 900, 901


def allocate_num_id(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(ids + [900]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None, color: str | None = None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(TOKENS.body_after_pt)
    p.paragraph_format.line_spacing = TOKENS.body_line_spacing
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=TOKENS.body_size_pt, bold=True, color=color or INK_BLUE)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=TOKENS.body_size_pt, color=color or "20242A")
    else:
        set_run_font(p.add_run(text), size=TOKENS.body_size_pt, color=color or "20242A")
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    num_id = allocate_num_id(doc, BULLET_ABSTRACT_ID)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), size=TOKENS.body_size_pt)


def add_steps(doc: Document, items: Iterable[str]):
    num_id = allocate_num_id(doc, DECIMAL_ABSTRACT_ID)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), size=TOKENS.body_size_pt)


def add_note(doc: Document, title: str, body: str, *, kind: str = "note"):
    colors = {
        "note": (CALLOUT, DARK_BLUE),
        "tip": ("EEF7F1", GREEN),
        "warning": ("FFF7E6", "7A5A00"),
        "risk": ("FDEEEE", RED),
    }
    fill, accent = colors.get(kind, colors["note"])
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(f"{title}  "), size=10.0, bold=True, color=accent)
    set_run_font(p.add_run(body), size=10.0, color="29313A")
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size="8")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(
            run,
            size={1: TOKENS.h1_size_pt, 2: TOKENS.h2_size_pt, 3: TOKENS.h3_size_pt}[level],
            bold=True,
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        )
    return p


def optimize_screenshot(source: Path) -> Path:
    OPTIMIZED_DIR.mkdir(parents=True, exist_ok=True)
    target = OPTIMIZED_DIR / f"{source.stem}.jpg"
    if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
        return target
    with Image.open(source) as im:
        rgb = ImageOps.exif_transpose(im).convert("RGB")
        if rgb.width > 1500:
            ratio = 1500 / rgb.width
            rgb = rgb.resize((1500, int(rgb.height * ratio)), Image.Resampling.LANCZOS)
        rgb.save(target, "JPEG", quality=88, optimize=True, progressive=True, subsampling=0)
    return target


def add_figure(doc: Document, source: Path, caption: str, *, width_in: float = 6.3):
    global FIGURE_COUNTER
    FIGURE_COUNTER += 1
    source = optimize_screenshot(source) if source.suffix.lower() != ".jpg" else source
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(source), width=Inches(width_in))
    inline._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cp.add_run(f"图 {FIGURE_COUNTER}  {caption}"), size=9.0, color=MUTED)
    return FIGURE_COUNTER


def find_font_path() -> Path:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError("No Chinese font found")


FONT_PATH = find_font_path()


def font(size: int, bold: bool = False):
    path = Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists() else FONT_PATH
    return ImageFont.truetype(str(path), size=size)


def draw_centered(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill: str):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=4, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text, font=fnt, fill=fill, spacing=4, align="center")


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#6B7A90", width=4):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    length = 13
    for delta in (2.55, -2.55):
        x = x2 + length * math.cos(angle + delta)
        y = y2 + length * math.sin(angle + delta)
        draw.line([(x2, y2), (x, y)], fill=color, width=width)


def create_flow_diagram(path: Path):
    im = Image.new("RGB", (1500, 430), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "标准建模工作流", font=font(38, True), fill="#0B2545")
    steps = [
        ("1", "选择方案/模型", "确认当前工作对象"),
        ("2", "放置元件", "从图元库拖入画布"),
        ("3", "连接端子", "建立能源网络拓扑"),
        ("4", "填写模型", "参数与量测保持一致"),
        ("5", "检查拓扑", "处理断点和告警"),
        ("6", "保存/导出", "持久化并形成交付文件"),
    ]
    x0, y0, w, h, gap = 48, 125, 205, 180, 35
    colors = ["#E8EEF5", "#EAF4F4", "#FFF4DD", "#F3EAF7", "#EEF7F1", "#EDEFF8"]
    for idx, (num, title, detail) in enumerate(steps):
        x = x0 + idx * (w + gap)
        d.rounded_rectangle((x, y0, x + w, y0 + h), radius=18, fill=colors[idx], outline="#9AA8B8", width=3)
        d.ellipse((x + 14, y0 + 14, x + 60, y0 + 60), fill="#2E74B5")
        draw_centered(d, (x + 14, y0 + 14, x + 60, y0 + 60), num, font(25, True), "white")
        draw_centered(d, (x + 14, y0 + 60, x + w - 14, y0 + 116), title, font(25, True), "#18324A")
        draw_centered(d, (x + 12, y0 + 112, x + w - 12, y0 + h - 10), detail, font(18), "#46566A")
        if idx < len(steps) - 1:
            arrow(d, (x + w + 4, y0 + h / 2), (x + w + gap - 5, y0 + h / 2))
    im.save(path, optimize=True)


def create_hierarchy_diagram(path: Path):
    im = Image.new("RGB", (1500, 650), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "类别库、类、派生类、元件与模型实例", font=font(38, True), fill="#0B2545")
    nodes = [
        ((70, 150, 320, 280), "类别库\n交流设备", "#E8EEF5"),
        ((390, 150, 640, 280), "基类\nACGenerator", "#DDECF8"),
        ((710, 80, 990, 215), "派生类\nACNuclearGen", "#F0E6F5"),
        ((710, 300, 990, 435), "直属元件\ncustom-ACGenerator", "#FFF1D6"),
        ((1060, 80, 1390, 215), "派生类元件\nac-nuclear-source", "#EAF4F4"),
        ((1060, 300, 1390, 435), "画布模型实例\ndev_type 可编辑", "#EEF7F1"),
    ]
    for box, text, fill in nodes:
        d.rounded_rectangle(box, radius=18, fill=fill, outline="#75869A", width=3)
        draw_centered(d, box, text, font(27, True), "#18324A")
    arrow(d, (320, 215), (390, 215))
    arrow(d, (640, 200), (710, 150))
    arrow(d, (640, 230), (710, 365))
    arrow(d, (990, 150), (1060, 150))
    arrow(d, (990, 365), (1060, 365))
    arrow(d, (1225, 215), (1225, 300))
    d.rounded_rectangle((125, 500, 1375, 600), radius=16, fill="#F7F9FC", outline="#BCC6D2", width=2)
    draw_centered(d, (145, 510, 1355, 590), "规则：基类自己的元件直接位于基类下；派生类单独作为二级节点，其元件位于派生类下。元件负责图形表现，类负责参数、量测和端子语义。", font(21), "#46566A")
    im.save(path, optimize=True)


def create_inheritance_diagram(path: Path):
    im = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "有效字段集合的形成方式", font=font(38, True), fill="#0B2545")
    boxes = [
        ((80, 150, 430, 440), "基类字段\nidx / name / run_stat\ndev_type / node\n额定值、上下限\n控制与设定字段", "#E8EEF5"),
        ((570, 150, 920, 440), "派生类新增字段\n只定义差异项\n不重复基类字段\n可增加专用控制量", "#F0E6F5"),
        ((1070, 150, 1420, 440), "模型实例有效集合\n基类 + 派生类\n去重、顺序稳定\n右侧模型表一致", "#EEF7F1"),
    ]
    for box, text, fill in boxes:
        d.rounded_rectangle(box, radius=20, fill=fill, outline="#75869A", width=3)
        draw_centered(d, box, text, font(25, True), "#18324A")
    arrow(d, (430, 295), (570, 295))
    arrow(d, (920, 295), (1070, 295))
    d.rounded_rectangle((175, 500, 1325, 580), radius=14, fill="#FFF7E6", outline="#C9A34A", width=2)
    draw_centered(d, (195, 510, 1305, 570), "派生类已有基类字段时，不应再次定义；元件不保存类参数和量测，只引用所属类的有效定义。", font(22), "#6A4D00")
    im.save(path, optimize=True)


def create_save_boundary_diagram(path: Path):
    im = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "三类保存边界", font=font(38, True), fill="#0B2545")
    columns = [
        (60, "保存类定义", ["参数字段", "量测定义", "基类端子能源属性", "继承差异项"], "#E8EEF5"),
        (520, "保存元件定义", ["图元图层", "图形样式", "端子显示位置", "中英文显示名"], "#FFF1D6"),
        (980, "保存模型", ["画布位置与连线", "实例参数值", "实例量测值", "图层与页面配置"], "#EEF7F1"),
    ]
    for x, title, items, fill in columns:
        d.rounded_rectangle((x, 130, x + 400, 500), radius=20, fill=fill, outline="#75869A", width=3)
        draw_centered(d, (x + 20, 145, x + 380, 215), title, font(29, True), "#18324A")
        y = 235
        for item in items:
            d.ellipse((x + 42, y + 8, x + 56, y + 22), fill="#2E74B5")
            d.text((x + 75, y), item, font=font(23), fill="#33465A")
            y += 58
    im.save(path, optimize=True)


def create_svg_flow_diagram(path: Path):
    im = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "SVG 图元交换与再编辑流程", font=font(38, True), fill="#0B2545")
    steps = [
        ("选择元件", "右键指定对象"),
        ("导出 SVG", "用户选择目录/文件名"),
        ("去除辅助对象", "不导出端子和连接线"),
        ("导入 SVG", "读取原始 viewBox 与几何"),
        ("生成可编辑图层", "支持增、删、改"),
        ("保存元件", "持久化图元显示"),
    ]
    x0, y0, w, h, gap = 55, 160, 205, 220, 35
    for idx, (title, detail) in enumerate(steps):
        x = x0 + idx * (w + gap)
        fill = ["#E8EEF5", "#EAF4F4", "#FFF1D6", "#F0E6F5", "#EEF7F1", "#EDEFF8"][idx]
        d.rounded_rectangle((x, y0, x + w, y0 + h), radius=18, fill=fill, outline="#75869A", width=3)
        draw_centered(d, (x + 12, y0 + 25, x + w - 12, y0 + 105), title, font(25, True), "#18324A")
        draw_centered(d, (x + 12, y0 + 105, x + w - 12, y0 + h - 18), detail, font(19), "#46566A")
        if idx < len(steps) - 1:
            arrow(d, (x + w + 3, y0 + h / 2), (x + w + gap - 5, y0 + h / 2))
    d.rounded_rectangle((150, 455, 1350, 560), radius=15, fill="#F7F9FC", outline="#BCC6D2", width=2)
    draw_centered(d, (170, 465, 1330, 550), "尺寸规则不依赖连接线是否存在；平台导出会记录源端子数量，旧 SVG 缺少元数据时按常规有端子设备处理。不会根据 dev-kind 套用内置模板。", font(21), "#46566A")
    im.save(path, optimize=True)


def create_export_directory_diagram(path: Path):
    im = Image.new("RGB", (1500, 540), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 30), "最近导出目录的复用机制", font=font(38, True), fill="#0B2545")
    nodes = [
        ((80, 150, 350, 350), "浏览器\nlocalStorage\n保存上次目录", "#E8EEF5"),
        ((450, 150, 720, 350), "发起导出\n仍然打开\n另存为窗口", "#FFF1D6"),
        ((820, 150, 1090, 350), "后台校验\n绝对路径且\n目录存在", "#F0E6F5"),
        ((1190, 150, 1460, 350), "用户选择\n新目录后\n覆盖旧记录", "#EEF7F1"),
    ]
    for box, text, fill in nodes:
        d.rounded_rectangle(box, radius=18, fill=fill, outline="#75869A", width=3)
        draw_centered(d, box, text, font(25, True), "#18324A")
    for i in range(3):
        arrow(d, (nodes[i][0][2] + 4, 250), (nodes[i + 1][0][0] - 4, 250))
    d.text((190, 420), "E、SVG 等导出类型共用同一最近目录；隐私模式无法写 localStorage 时，文件导出仍可正常完成。", font=font(22), fill="#46566A")
    im.save(path, optimize=True)


def create_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    creators = {
        "workflow.png": create_flow_diagram,
        "hierarchy.png": create_hierarchy_diagram,
        "inheritance.png": create_inheritance_diagram,
        "save-boundaries.png": create_save_boundary_diagram,
        "svg-flow.png": create_svg_flow_diagram,
        "export-directory.png": create_export_directory_diagram,
    }
    for name, creator in creators.items():
        creator(DIAGRAM_DIR / name)


def add_section(doc: Document, title: str, paragraphs: Sequence[str], *, bullets: Sequence[str] | None = None, note: tuple[str, str, str] | None = None, level: int = 2):
    add_heading(doc, title, level)
    for paragraph in paragraphs:
        add_para(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)
    if note:
        add_note(doc, note[0], note[1], kind=note[2])


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("用户操作与维护手册"), size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("电力能源系统图上建模平台"), size=28, bold=True, color=INK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("从方案管理、画布建模到类/元件定义与文件交换"), size=14, color=DARK_BLUE)

    add_table(
        doc,
        ["版本", "适用程序", "发布日期", "文档定位"],
        [["V2.0 紧凑图文版", "main@1267a934", "2026-08-17", "操作、配置、维护与排障"]],
        [1800, 2300, 1700, 3560],
    )
    add_figure(doc, ASSET_DIR / "02-model-ieee14.png", "平台示例模型与三栏式工作界面", width_in=6.25)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("适用于建模人员、设备模型维护人员、接口配置人员和平台管理员"), size=10, italic=True, color=MUTED)
    doc.add_page_break()


def build_manual() -> Document:
    doc = Document()
    configure_document(doc)
    global BULLET_ABSTRACT_ID, DECIMAL_ABSTRACT_ID
    BULLET_ABSTRACT_ID, DECIMAL_ABSTRACT_ID = create_numbering(doc)
    add_cover(doc)

    add_heading(doc, "文档说明", 1)
    add_para(doc, "本手册以当前运行版本为依据，系统说明电力能源系统图上建模平台的界面结构、数据层级、标准操作流程、保存边界和常见问题。文档既面向首次使用者，也面向需要维护类参数、量测、端子和图元样式的高级用户。")
    add_para(doc, "平台采用“类别库—类—派生类—元件—模型实例”的多层结构。理解这套结构，是避免字段重复、元件保存越界、量测绑定失效以及 SVG 图形覆盖类定义的关键。本手册在每个操作章节中都明确说明“当前编辑对象是谁”“保存后写到哪里”“哪些内容由上层继承”。")
    add_note(doc, "阅读建议", "新用户先阅读第 1～4 章并完成标准工作流；负责设备库维护的人员重点阅读第 8～13 章；负责数据交换和交付的人员重点阅读第 14～16 章。", kind="tip")

    add_heading(doc, "内容索引", 2)
    index_rows = [
        ("1～4", "平台概览、术语、环境、主界面和快速上手"),
        ("5～7", "方案/模型、画布、图形属性、拓扑和量测"),
        ("8～10", "图元库层级、新建类与元件、字段继承"),
        ("11～13", "类定义、元件图元编辑、复制粘贴和 SVG"),
        ("14～16", "模板库、E 文件、导出目录、持久化和交付"),
        ("17", "故障定位和问题处理"),
        ("附录 A～C", "字段速查、验收清单和维护边界"),
    ]
    add_table(doc, ["章节", "主要内容"], index_rows, [1800, 7560], caption="手册内容索引")

    add_heading(doc, "1 平台定位与核心术语", 1)
    add_section(doc, "1.1 平台解决什么问题", [
        "平台将电力、直流、氢能和热能设备放在统一画布中，通过可视化拖拽、端子连接、模型参数、动态量测和文件接口完成系统建模。它既是图形编辑器，也是设备类型管理器和模型数据维护工具。",
        "与普通绘图软件不同，画布上的每个设备都具有明确的类定义和元件显示定义。图形只负责“看起来是什么”，类负责“具有什么字段和量测”，模型实例负责“在当前工程中的具体取值和连接关系”。",
    ], bullets=[
        "方案用于组织一组相关模型；模型对应一个可独立保存和导出的页面或系统。",
        "类别库用于按能源或业务类别组织类；类用于定义字段、量测和端子语义。",
        "元件是可拖入画布的具体图形类型，可具有不同中文名、英文名和图元样式。",
        "模型实例是元件落到画布后的实际对象，保存位置、连接线、参数值和量测值。",
    ])
    add_section(doc, "1.2 类、派生类和元件的区别", [
        "基类定义通用字段。例如 ACGenerator 统一定义 idx、name、run_stat、dev_type、node、控制类型和设定值。派生类只增加差异字段，不重复基类已经存在的字段。这样能够保证同一类设备的有效参数集合稳定，避免模型表出现重复列或字段丢失。",
        "元件不拥有独立的类参数表和量测表。元件只保存图元显示信息、端子显示位置和可修改的中英文显示名。一个基类或派生类可以对应多个元件，多个元件共享同一套类语义，但可以采用不同图形。",
    ], note=("关键原则", "如果基类已经有某字段，派生类不再重复定义；如果当前选中对象是元件，界面只显示图元定义，不显示类参数和量测定义。", "warning"))
    add_figure(doc, DIAGRAM_DIR / "hierarchy.png", "类别库、基类、派生类、元件与模型实例的层级关系")

    add_heading(doc, "2 运行环境与进入平台", 1)
    add_section(doc, "2.1 服务与访问", [
        "平台由前端 WEB 和后台服务组成。标准开发运行环境中，前端通过 5173 端口提供 Vite 页面，后台通过 5174 端口提供模型库、元件库、导入导出和持久化接口。浏览器访问前端地址后，页面会自动建立模型库连接。",
        "判断系统是否可用不能只看端口是否监听。至少应确认主页面能够返回、模型库能够加载、后台 /swigger 页面可访问，并能在左栏看到方案和模型记录。首次加载时出现“正在连接模型库”通常是数据初始化过程，可等待片刻；长期不恢复则按第 17 章检查后台。",
    ])
    add_section(doc, "2.2 首次进入后的检查", [
        "进入页面后先看顶部当前模型标识、左侧模型库、右侧属性栏和底部运行状态。若未选择模型，画布为空、保存按钮禁用属于正常状态。双击模型后，顶部应显示“方案 / 模型”，底部出现元件和联络线数量，画布呈现设备网络。",
    ], bullets=[
        "确认左侧存在“模型库、图元库、模板库”三个标签。",
        "确认顶部工具栏包含保存、E 文件接口定义、导出文件、拓扑和显示控制。",
        "确认右侧能够在基础、图元树和图元属性之间切换。",
        "确认底部能够看到坐标、缩放、拓扑、告警、日志和对象数量。",
    ])
    add_figure(doc, ASSET_DIR / "01-main-interface.png", "未打开模型时的主界面与功能区")

    add_heading(doc, "3 快速上手与标准工作流", 1)
    add_section(doc, "3.1 六步完成一个模型", [
        "标准工作流从明确工作对象开始，以保存和导出结束。建议不要在未确认当前方案/模型时拖入设备，也不要在拓扑告警未处理时直接交付文件。平台的保存按钮只保存当前模型；类定义和元件定义分别在定义窗口中保存。",
    ])
    add_figure(doc, DIAGRAM_DIR / "workflow.png", "从选择模型到保存导出的六步工作流")
    add_steps(doc, [
        "在模型库中双击目标模型，确认顶部当前模型名称正确。",
        "切换到图元库，展开类别、类或派生类，将元件拖入画布。",
        "从设备端子拖动到兼容端子，形成连接线；必要时调整折点和走向。",
        "选中设备，在右侧模型页填写实例参数，并在量测页确认动态量测。",
        "执行图上拓扑，查看底部拓扑状态和告警窗口，修正悬空端子或不兼容连接。",
        "点击保存写入当前模型；按需要导出 E、SVG 或其他交付文件。",
    ])
    add_section(doc, "3.2 保存前的最小检查", [
        "保存前至少确认当前模型、对象数量、拓扑状态和右侧字段完整性。若保存按钮处于禁用状态，表示当前模型没有新的修改；类参数、类量测或元件图元发生修改时，应在对应定义窗口保存，而不是依赖主界面保存按钮。",
    ], note=("避免误区", "“保存元件定义”不会保存类参数和量测；“保存类定义”不会保存元件图形；“保存”不会代替前两者。", "warning"))

    add_heading(doc, "4 主界面与工作区", 1)
    add_section(doc, "4.1 三栏式结构", [
        "主界面由左侧资源栏、中部画布和右侧属性栏组成。左侧负责选择模型、图元和模板，中部负责摆放与连接，右侧负责查看和修改当前选中对象。左右栏都支持永久显示、自动显示/隐藏和永久隐藏，可在小屏幕上释放画布空间。",
        "顶部工具栏提供全局操作，底部状态栏提供实时反馈。用户在任何操作前都应先确认顶部当前模型；在拖拽、连接或批量选择后，应观察底部对象数量、选中数和日志。",
    ])
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-01.jpg", "主界面功能区划分和各区域职责")
    add_section(doc, "4.2 顶部工具栏", [
        "图层管理用于激活和控制模型图层；编辑/浏览模式决定画布是否可修改；对齐标线帮助设备移动时保持整齐；图上拓扑用于计算并标记连接关系。颜色切换、配色设置和电压等级设置控制显示方案，不直接改变电气参数。",
        "分类图标库用于维护可复用图标，用户自定义修改管理用于查看内置定义的覆盖修改。组合、显示层级、对齐和旋转操作只有在选择了适用对象后才启用。禁用状态是权限和上下文约束的一部分，不应通过绕过界面直接触发。",
    ])
    add_section(doc, "4.3 视口与小地图", [
        "适配视图会把当前模型缩放到可见区域；居中选中和缩放到选中区域要求先选中设备或连接线；放大、缩小和重置缩放用于精细调整。小地图提供全局位置参考，收紧画布根据已有对象边界压缩有效画布。",
    ], bullets=[
        "大模型先使用适配视图，再通过小地图定位局部区域。",
        "编辑图元细节时避免过低缩放，否则端子和选择框难以准确操作。",
        "批量移动后检查画布边界，防止对象被自动扩界后留在远处。",
    ])

    add_heading(doc, "5 方案和模型库管理", 1)
    add_section(doc, "5.1 选择、打开与定位模型", [
        "左侧模型库按“方案—模型”组织。单击记录用于选择并在右侧查看基础属性；双击模型才会真正加载画布。顶部当前模型、左下方案 ID 和模型 ID 是判断是否打开成功的权威信号。",
        "搜索框可按方案或模型名称过滤。对大型模型库，先搜索方案，再双击其中模型，比在长列表中滚动更可靠。切换模型前若当前模型有未保存修改，平台会提示保存、放弃或继续编辑。",
    ])
    add_figure(doc, ASSET_DIR / "02-model-ieee14.png", "双击 IEEE14 后加载模型、背景页面和右侧基础属性")
    add_section(doc, "5.2 模型基础属性", [
        "右侧基础页可编辑模型名称、显示宽度、显示高度、背景色、背景图、背景页面、背景图层、功率/电压/电流单位、功率基值、区域和模型类型。显示尺寸影响画布范围，不等同于导出图片分辨率。",
        "背景页面适合在主接线图上叠加专题页面。选择背景页面后，还可指定参与显示的背景图层。避免把当前模型本身选为背景页面，以免形成循环引用；若背景页面内容较复杂，应检查画布性能和导出结果。",
    ])
    add_section(doc, "5.3 右键操作与数据安全", [
        "方案和模型列表支持右键操作。执行新建、复制、重命名、导入、导出或删除前，先确认当前选中层级。删除属于不可逆业务操作，应确认目标 ID 和依赖关系；复制适合快速建立相似页面，但复制后必须修改名称并检查背景页面、图层和模型参数。",
    ])
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-03.jpg", "模型库右键菜单与上下文操作")

    add_heading(doc, "6 画布建模与连接", 1)
    add_section(doc, "6.1 放置元件", [
        "切换到图元库后，按类别、基类和派生类找到目标元件。元件标签前的小图标是图元缩略图，可与派生类标签区分。把元件拖到画布后，平台生成模型实例，并以所属类英文名作为 dev_type 默认值；dev_type 可以在模型表中修改。",
        "基类自己的元件直接位于基类下；派生类元件位于派生类二级节点下。若某个类下没有元件，仅能维护类定义，不能直接拖入画布。元件数量显示在类标签右侧，可用于快速判断展开后是否应出现可用图元。",
    ])
    add_figure(doc, ASSET_DIR / "03-component-library.png", "图元库、能源类别和类列表")
    add_figure(doc, ASSET_DIR / "04-class-tree-expanded.png", "向右浮动展开交流电源类及其元件")
    add_section(doc, "6.2 端子连接", [
        "端子是模型拓扑的连接点。拖动连接时，平台根据端子能源属性和角色判断是否兼容。基类定义端子能源属性，派生类继承，元件只保存端子在图形上的显示位置。多端设备应逐个确认端子编号与业务语义。",
        "连接线可以具有直线、折线或自适应路径。移动设备时连接线随端子更新；如果使用手动折点，移动后应检查折点是否仍合理。连接完成不代表拓扑一定正确，仍需执行图上拓扑检查节点归并和悬空端子。",
    ], bullets=[
        "交流端子只连接交流网络，直流、氢能和热能端子遵循相同原则。",
        "端子显示位置变化属于元件定义；端子能源属性变化属于基类定义。",
        "导出单个元件 SVG 时端子和端子连接线不会导出，但模型连接关系不受影响。",
    ])
    add_section(doc, "6.3 选择、移动和组合", [
        "单击选择一个对象，按住 Ctrl 可添加或移除多选对象，框选可一次选择区域内对象。组合用于把多个图形作为整体操作；解除组合恢复独立对象。对齐操作可按左、右、上、下、水平中心或垂直中心对齐，旋转操作用于批量调整方向。",
    ])

    add_heading(doc, "7 图形属性、模型参数和动态量测", 1)
    add_section(doc, "7.1 图形属性", [
        "选中元件后，右侧图形页显示位置、尺寸、旋转、字体、线型、颜色、显示层级和标签等外观属性。图形属性只影响当前模型实例的显示；若希望修改图元库中的默认样式，应进入元件定义窗口编辑元件图元。",
        "内置图元的样式也允许修改。修改后会产生自定义覆盖并持久化到后台，进程重启后仍保留。需要恢复系统默认样式时，使用用户自定义修改管理中的恢复功能，而不是手工猜测原始参数。",
    ])
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-05.jpg", "选中设备后的图形属性面板", width_in=5.65)
    add_heading(doc, "7.2 模型参数", 2)
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-06.jpg", "右侧模型表展示设备有效参数", width_in=5.65)
    add_para(doc, "模型页使用与类定义窗口完全相同的有效参数集合。有效集合由基类字段和派生类新增字段合并、去重并保持稳定顺序。这样可以保证 ACGenerator 等类在定义窗口、右侧模型表和文件导出中看到一致字段。")
    add_para(doc, "dev_type 默认填写元件所属类（基类或派生类）的英文名称，而不是元件英文名；用户可以修改。拓扑相关 node、i_node、j_node 等字段由端子定义自动形成，通常通过连接关系更新，不应把它们误当成普通显示字段删除。")
    add_heading(doc, "7.3 动态量测", 2)
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-07.jpg", "右侧量测页和量测值编辑", width_in=5.65)
    add_para(doc, "量测页显示所属类定义的量测项，并把量测绑定到有效参数字段。常见量测包括有功、无功、电压、电流、频率、状态、压力、温度、流量、SOC 和储量。量测默认显示策略可跟随类型，也可强制显示或隐藏。")
    add_para(doc, "出现“没有绑定量测”时，先检查选中元件所属类是否正确，再检查类量测中的关联字段是否属于有效参数集合。派生类应继承基类量测；元件本身不单独保存量测。")
    add_section(doc, "7.4 图上拓扑与告警", [
        "点击图上拓扑后，平台根据端子连接、节点字段和设备角色建立网络关系。底部状态栏显示拓扑状态，告警窗口列出悬空、冲突或无法归并的对象。拓扑成功后再导出模型，可以降低接口文件中节点缺失的风险。",
    ], note=("排障顺序", "先处理悬空端子，再处理能源类型不匹配，最后检查 node/topo 字段；不要直接手工修改大量节点号掩盖连接问题。", "warning"))

    add_heading(doc, "8 图元库、类层级和搜索", 1)
    add_section(doc, "8.1 图元库的展开方式", [
        "图元库支持向下展开和向右浮动两种方式。向下展开适合逐级浏览；向右浮动可在不拉长左栏的情况下显示类下元件，适合大类。菜单和浮层会相对鼠标及触发项定位，点击其他标签或空白区域后应自动关闭。",
        "搜索框同时匹配类别库、类、派生类和元件的中英文名称。搜索后仍保留层级关系，便于判断结果属于基类、派生类还是具体元件。",
    ])
    add_section(doc, "8.2 二级派生列表", [
        "存在派生类时，树采用二级列表：基类下先列派生类，再在派生类下列相关元件；基类自己的元件直接列在基类下。该规则同时适用于左侧图元库和元件定义窗口的结构树。",
        "元件标签带缩略图和“内置/自定义”标识，派生类标签显示字段层级但不显示元件缩略图。不要仅凭中文名称判断对象类型，英文类名和树层级才是可靠依据。",
    ])
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-08.jpg", "类别库、基类、派生类与元件关系示意")
    add_section(doc, "8.3 类导入导出", [
        "图元库顶部的导入和导出用于交换类及相关元件定义。导入前应检查英文名冲突和继承目标，导入后检查参数、量测、端子和图元是否完整。类导出与单个元件 SVG 导出用途不同：前者用于迁移结构化定义，后者只用于交换图形。",
    ])

    add_heading(doc, "9 新建类别、类、派生类和元件", 1)
    add_section(doc, "9.1 创建权限与上下文", [
        "创建入口位于元件定义窗口左侧结构树的右键菜单。对类别库、基类、派生类和元件右键时，可用菜单项不同：类别库上下文允许创建基类；类上下文允许创建派生类和元件；元件上下文不允许继续创建子级。界面禁用状态和确认逻辑同时执行校验。",
        "新建类别、类或元件时，应先单击目标树节点再右键。菜单出现后如果切换标签或单击其他对象，菜单会自动消失。菜单位置以当前鼠标为准，不应偏离触发点。",
    ])
    add_section(doc, "9.2 新建类", [
        "新建类确认后，平台立即把类基础记录写入后台并刷新左侧列表，然后在右侧进入该类定义界面。这样可以避免用户在尚未落库的临时对象上编辑大量内容。类中文名用于显示，英文名用于稳定标识和继承引用，应在同一类别库中保持唯一。",
        "基类创建后自动具备公共字段，如 idx、name、run_stat、dev_type；根据端子数量和能源属性生成 node、i_node、j_node 或其他 topo 相关字段。派生类自动继承基类字段，只需定义差异项。",
    ], note=("英文名规则", "英文类名应使用稳定的英文标识，创建后不要通过普通字符串替换修改 kind；需要调整显示时修改中文名或英文显示名。", "warning"))
    add_section(doc, "9.3 新建元件", [
        "新建元件弹窗要求填写中文名和英文名。确认后，元件记录立即写入后台并出现在左侧树中，右侧整个页面进入元件模式：顶部显示元件名称，中部是图元编辑器，底部保存按钮保存元件图形。",
        "元件模式不显示设备参数定义和量测定义按钮。保存元件时只保存图元显示信息、状态图层和端子显示位置，不保存所属类的参数、量测或端子能源属性。中文名和英文名都可以在元件定义页直接修改，左上角不再提供单独重命名按钮。",
    ])

    add_heading(doc, "10 字段继承、默认字段和有效参数集合", 1)
    add_section(doc, "10.1 基类默认字段", [
        "新建基类会生成系统公共字段。idx 和 name 用于实例标识，run_stat 表示工作状态，dev_type 表示设备类型；端子字段由端子定义生成。不同设备类还会补充额定值、上下限、控制方式和设定值。",
        "以 ACGenerator 为例，有效字段应包括控制类型 control_type、有功设定值 p_set、无功设定值 q_set、电压设定值 v_set、调节系数 alpha 等；交流电源 regable 使用数字枚举类型。特定控制设备还可包含 i_p_set、j_p_set、i_i_set、j_i_set、i_v_set、j_v_set 等端侧设定字段。",
    ])
    add_table(doc, ["字段组", "典型字段", "说明"], [
        ["公共标识", "idx, name, rdf_id", "实例编号、名称和外部原始 ID"],
        ["运行与类型", "status, run_stat, dev_type", "状态、投停和设备类英文名"],
        ["拓扑", "node / i_node / j_node", "按端子定义自动生成并随连接更新"],
        ["额定与限值", "rated_voltage, p_max, q_min", "额定量及运行上下限"],
        ["控制", "control_type, regable, alpha", "控制模式、可调性和调节系数"],
        ["设定值", "p_set, q_set, v_set", "控制目标值，可根据类扩展"],
    ], [1600, 3000, 4760], caption="类字段的常见分组")
    add_section(doc, "10.2 派生类去重", [
        "派生类读取基类有效字段后，只保存新增或覆盖所需的差异定义。如果派生类表中存在与基类同名字段，平台会把它识别为重复，不应依赖重复行改变字段顺序。推荐在基类集中维护公共字段，在派生类维护技术特性。",
    ])
    add_figure(doc, DIAGRAM_DIR / "inheritance.png", "基类、派生类和模型实例有效字段的合并方式")
    add_section(doc, "10.3 dev_type 的默认值", [
        "设备实例的 dev_type 默认值是元件所属类的英文名：若元件直属基类，则使用基类英文名；若元件位于派生类下，则使用派生类英文名。该值可以在模型实例中编辑，用于外部接口映射或业务分类。",
        "dev_type 不是元件英文名。元件英文名例如 ac-wind-source，主要用于图元类型和显示模板识别；所属类英文名例如 ACWindGen，表示参数和量测语义来源。两者应分别维护。",
    ], note=("注意", "若外部模型协议明确要求 dev_type 使用元件英文名，应通过接口映射实现，不要破坏平台默认的类语义。", "note"))

    add_heading(doc, "11 类参数、量测和端子定义", 1)
    add_section(doc, "11.1 打开类定义", [
        "在图元库选择类，点击“元件定义”进入统一的类/元件管理窗口，再在左侧树选中目标类。右侧顶部显示类别库、所属类、派生关系、端子数量和容器属性。基类还显示端子能源属性；派生类不重复配置。",
        "参数定义和量测定义都具有编辑标识。发生增、删、改后，底部状态会提示存在未保存修改；切换类、切换标签或关闭窗口时，会询问保存、放弃或继续编辑。",
    ])
    add_figure(doc, ASSET_DIR / "05-class-definition.png", "ACGenerator 基类的参数定义和端子能源属性", width_in=5.65)
    add_section(doc, "11.2 参数表操作", [
        "参数表支持新增、复制、上移、下移、删除和行选择。中文名称面向用户，英文名称是稳定字段键；取值类型支持整数、浮点数、字符串、字符串枚举和数字枚举。枚举字段可编辑枚举项和默认值。",
        "系统生成字段如 idx、name、run_stat、dev_type 和 topo 字段可能以只读方式出现，防止误删破坏模型契约。普通业务字段可直接编辑。调整顺序会影响类定义窗口、右侧模型表和导出文件的展示顺序。",
    ], bullets=[
        "新增前先检查英文名是否已经由基类提供。",
        "复制字段后必须修改英文名，避免重复键。",
        "数字枚举项应同时维护数值和中文含义。",
        "默认值为空表示新实例不强制填充，不等同于数值 0。",
    ])
    add_heading(doc, "11.3 量测表操作", 2)
    add_figure(doc, ASSET_DIR / "06-class-measurements.png", "ACGenerator 类量测定义及关联字段", width_in=5.65)
    add_para(doc, "量测定义支持添加、复制、排序和删除。每条量测包含名称、类型、位置、关联字段和默认显示策略。关联字段下拉框来自当前类有效参数集合，因此 control_type、p_set、q_set、v_set 等字段只有在有效集合中存在时才能被绑定。")
    add_para(doc, "量测位置可指向设备本体或端子。多端设备应把端侧量测放在正确端子，避免 i、j 侧数据颠倒。默认显示选择“跟随类型”时，由系统量测类型配置决定；选择“显示/隐藏”会覆盖该默认。")
    add_section(doc, "11.4 端子能源属性", [
        "端子能源属性从元件处上移到基类定义。基类按端子编号选择交流、直流、氢能或热能；派生类继承，不显示重复控件。元件定义只调整端子图形位置，不改变能源语义。",
        "修改端子数量会影响自动生成的 topo 字段和所有元件的端子显示。此操作风险较高，应先确认已有模型实例和外部接口的兼容性，再更新对应元件图元。",
    ])

    add_heading(doc, "12 元件图元编辑与未保存状态", 1)
    add_section(doc, "12.1 元件定义页面", [
        "在结构树中选择元件后，右侧切换为元件模式。顶部可修改元件中文名和英文名，查看所属类、派生关系和端子数量；中部是图元编辑器；底部保存按钮写入当前元件。类参数、类量测和端子能源属性不会出现在元件页面。",
        "图元编辑器支持基础元素、静态图元库、内置图标、分类图标和外部图标。图形可由线、折线、点、矩形、圆、椭圆、圆弧、文本框等多个图层组合。每个状态拥有独立图层集合，可新增运行状态图案。",
    ])
    add_figure(doc, ASSET_DIR / "07-component-editor.png", "交流核能发电机的可编辑图元与端子位置")
    add_section(doc, "12.2 图层的增、删、改", [
        "单击图层后切换到“选中图元”页，可修改位置、尺寸、旋转、颜色、线宽、填充和文本。删除只删除当前图层；新增元素插入当前状态。拖动或修改属性后出现编辑标识，表示尚未保存。",
        "粘贴生成的元件、内置原始元件和普通 SVG 导入形成的元件都支持继续编辑，不会作为不可拆分的背景图片锁死。对复杂 SVG，平台把可识别几何转换为图层；无法转换的部分仍以可移动 SVG 图层保留。",
    ])
    add_section(doc, "12.3 文字、端子和状态", [
        "文本框图层可以编辑内容、字体、字号、颜色、对齐和旋转。若单击文字后无法输入，应确认已选中“选中图元”属性页，而不是全局信息页。端子可在预览区拖动，也可输入归一化 X/Y 位置。",
        "端子位置应尽量靠近图形外缘，连接线从端子向外延伸。图元导出为 SVG 时不包含端子和端子连接线；端子显示仍保存在平台元件定义中。",
    ])
    add_figure(doc, ASSET_DIR / "legacy" / "legacy-11.jpg", "文本图层编辑和图元属性")
    add_section(doc, "12.4 编辑标识与关闭提示", [
        "图元样式、类参数和类量测发生修改后都会进入脏状态。切换页面、选择其他类/元件或关闭窗口时，平台提供保存后继续、不保存继续和继续编辑。继续编辑返回当前对象；保存后继续先持久化；不保存继续丢弃本轮草稿。",
    ], note=("操作建议", "看到编辑标识时先确认修改对象和保存边界；不要通过刷新浏览器规避提示，否则草稿可能丢失。", "warning"))

    add_heading(doc, "13 元件复制粘贴与 SVG 导入导出", 1)
    add_section(doc, "13.1 右键菜单", [
        "元件结构树的右键菜单包含新建类别、新建类、新建元件、复制元件、粘贴为新元件、导出图元为 SVG、导入 SVG 为图元和删除。菜单项是否启用由当前对象类型和复制缓冲区决定。",
        "复制只能针对元件；复制成功后，粘贴可在同一类或其他类上使用，并支持一次复制、多次粘贴。切换标签或单击其他区域后菜单自动关闭，下一次右键会在新鼠标位置重新显示。",
    ])
    add_figure(doc, ASSET_DIR / "08-component-context-menu.png", "元件右键菜单及复制、粘贴、SVG 操作")
    add_section(doc, "13.2 复制与粘贴", [
        "复制操作把元件图元显示、状态图层、端子显示位置和可变形属性写入前端复制缓冲区。粘贴时弹出名称窗口，要求设置新中文名和英文名。系统先执行名称查重，通过后立即写入后台并刷新树，不要求用户再点击一次保存才能看到结果。",
        "目标类可以与源类相同，也可以不同。跨类粘贴后，新元件使用目标类的参数、量测和端子语义，但保留源元件图形。若目标类端子数量不同，应进入图元编辑器检查端子显示位置。",
    ], bullets=[
        "粘贴菜单禁用时，确认复制动作发生在元件而不是类节点上。",
        "同一复制缓冲区可连续粘贴，直至复制其他元件或刷新页面。",
        "英文名必须唯一；中文名建议体现业务用途和所属类型。",
        "粘贴成功后仍可继续增、删、改图元图层。",
    ])
    add_section(doc, "13.3 导出 SVG", [
        "右键元件选择“导出图元为 SVG”后，平台弹出另存为窗口，允许指定目录和文件名。导出内容只包含元件主体图形，不包含端子和端子连接线，以便在外部矢量工具中作为纯图形使用。",
        "新导出的平台 SVG 在设备 use 元素上记录源端子数量，用于再次导入时恢复正确编辑画布尺寸。这个元数据不改变可见图形，也不把端子画进 SVG。",
    ])
    add_section(doc, "13.4 导入 SVG", [
        "选择“导入 SVG 为图元”后指定 SVG 文件。平台读取 viewBox、use、symbol 和普通几何，将图形放入当前元件编辑器。普通 SVG、平台导出 SVG 和包含多图层的 SVG 都应保持可编辑。",
        "导入尺寸不再依据 SVG 中是否存在端子连接线判断。对于平台新 SVG，使用源端子数量选择 180×120 或 240×160 编辑框；旧平台 SVG 缺少元数据时按常规有端子设备处理。平台不会根据 dev-kind 回查内置模板，也不会用模板几何替换用户文件。",
    ], note=("尺寸检查", "同一个主体几何无论是否带端子连接线，导入后的主体大小应一致。若仍异常，比较外层 viewBox、use 尺寸、symbol viewBox 和 data-export-source-terminal-count。", "tip"))
    add_figure(doc, DIAGRAM_DIR / "svg-flow.png", "SVG 导出、导入、可编辑化和保存流程")

    add_heading(doc, "14 模板库、E 文件接口与文件交换", 1)
    add_heading(doc, "14.1 模板库", 2)
    add_figure(doc, ASSET_DIR / "09-template-library.png", "模板库标签及导入导出入口", width_in=5.65)
    add_para(doc, "模板库用于保存可复用的局部模型或组合结构。模板可以包含多个元件及连接关系，适合重复放置标准站型、储能单元或多能耦合模块。模板库同样支持搜索、向下展开、向右浮动以及整体导入导出。")
    add_para(doc, "模板与元件的区别在于：元件对应单个设备图形和一个类语义，模板对应多个模型实例的组合。修改模板不会自动修改已经放入模型的实例；需要更新时应重新应用模板或手工同步。")
    add_section(doc, "14.2 E 文件接口定义", [
        "E 文件接口定义用于配置模型类、字段和外部文件列之间的映射、顺序与输出规则。类字段的有效集合会进入接口定义，重复或缺失字段会影响导入导出。维护接口前先确保基类和派生类字段已保存。",
        "导出 E 文件前检查模型拓扑、单位、功率基值和字段默认值。导入 E 文件后检查未识别类型、节点映射、设备数量和量测绑定，不要仅以文件解析成功作为建模正确的依据。",
    ])
    add_section(doc, "14.3 导出文件的一般规则", [
        "顶部“导出文件”汇总模型级导出功能。不同文件类型可能采用不同扩展名和内容生成器，但本地 WEB 环境下都优先使用系统另存为窗口。用户可以修改建议文件名，也可以选择新目录。",
    ])

    add_heading(doc, "15 最近导出目录、浏览器存储和持久化", 1)
    add_section(doc, "15.1 最近导出目录", [
        "平台在浏览器 localStorage 中保存最近一次本地导出目录，键名为 graph-modeling-platform.native-export.directory。E、SVG 等文件类型共享这一目录记录。再次导出时仍然弹出目录和文件名选择窗口，只是默认打开上次目录。",
        "后台接收 initialDirectory 后校验它是绝对路径且目录存在；无效或已删除的目录不会直接使用，而是回退到该导出功能的默认位置。用户本次选择的新目录会返回浏览器并覆盖旧记录。",
    ])
    add_figure(doc, DIAGRAM_DIR / "export-directory.png", "浏览器本地记录和系统另存为窗口之间的目录复用")
    add_section(doc, "15.2 浏览器限制与回退", [
        "在隐私模式、受限浏览器配置或 localStorage 不可用时，目录记忆可能失效，但导出仍继续。若本地原生对话框不可用，平台可回退到浏览器下载机制；浏览器出于安全原因通常不会暴露真实目录路径，因此无法完全复现原生默认目录。",
    ])
    add_section(doc, "15.3 内置图元修改持久化", [
        "对内置图元的样式修改会作为用户自定义覆盖写入后台，而不是只留在当前前端内存。WEB 进程重启后，平台重新加载覆盖记录并应用到内置图元。用户自定义修改管理提供查看、恢复和清理入口。",
        "持久化覆盖与程序内置默认值分离，既能保留用户修改，也能在需要时恢复系统版本。升级程序时应备份覆盖数据，并在升级后抽查常用内置元件。",
    ])

    add_heading(doc, "16 常用业务流程与交付", 1)
    add_section(doc, "16.1 新建派生类并建立元件", [
        "这个流程适合在 ACGenerator 等基类上增加风电、光伏、核电或储能特性。先创建派生类并维护差异字段，再创建或粘贴元件图形，最后在模型中实例化验证。",
    ])
    add_steps(doc, [
        "在元件定义结构树中选择基类，右键新建类，填写派生类中英文名。",
        "确认后等待后台写入和树刷新，选中新派生类。",
        "检查继承字段，不重复创建 idx、name、run_stat、dev_type 和基类控制字段。",
        "新增派生类专用字段和量测，保存类定义。",
        "右键派生类新建元件，或复制已有元件后粘贴到派生类。",
        "编辑元件图形和端子显示位置，保存元件定义。",
        "拖入测试模型，确认右侧模型表、量测和 dev_type。",
    ])
    add_section(doc, "16.2 跨类复制图元", [
        "当不同类需要相同或相似图形时，不必重画。复制源元件，在目标类上粘贴为新元件，设置唯一中英文名，检查目标端子数量后保存。复制只复用图形，目标类的参数和量测不会被源类覆盖。",
    ])
    add_section(doc, "16.3 从 SVG 建立新元件", [
        "先在目标类下新建元件，使基础记录入库；再在元件右键菜单选择导入 SVG。导入后检查主体大小、图层可编辑性和端子位置，必要时添加文本或状态图层，最后保存元件定义。",
    ])
    add_section(doc, "16.4 模型交付", [
        "模型交付前执行拓扑检查，保存当前模型，核对字段和量测，按目标系统配置 E 接口，再导出文件。导出窗口默认打开最近目录，但仍应确认本次文件名和目录。交付包建议同时包含接口说明、类版本和模型截图。",
    ])
    add_figure(doc, DIAGRAM_DIR / "save-boundaries.png", "类定义、元件定义和模型保存的内容边界")
    add_table(doc, ["检查项", "通过标准", "发现问题后的处理"], [
        ["当前模型", "顶部方案/模型与交付目标一致", "返回模型库重新双击"],
        ["拓扑", "无悬空和能源冲突告警", "检查端子与连接线"],
        ["有效字段", "右侧模型表与类定义一致", "检查继承、去重和保存"],
        ["量测", "量测存在且关联字段有效", "回到类量测定义修正"],
        ["图元", "图形、端子和缩放显示正常", "进入元件定义编辑并保存"],
        ["导出", "文件名、目录、格式和内容正确", "重新打开另存为窗口"],
    ], [1800, 3500, 4060], caption="模型交付前检查表")

    add_heading(doc, "17 常见问题与排查方法", 1)
    troubleshooting = [
        ("模型单击后画布仍为空", "单击只选择记录，双击模型后才加载。检查顶部当前模型和底部对象数量。"),
        ("模型表字段少于类定义", "检查基类/派生类有效集合、重复字段去重和类定义是否已保存。"),
        ("派生类出现大量重复字段", "删除派生类中基类已有字段，只保留差异项。"),
        ("提示没有绑定量测", "检查元件所属类、量测关联字段以及字段是否属于有效参数集合。"),
        ("粘贴菜单始终禁用", "确认右键复制的是元件节点；类节点不能写入元件复制缓冲区。"),
        ("粘贴后没有原图元", "检查复制缓冲区是否包含状态图层，粘贴后选择新元件确认图元编辑器。"),
        ("SVG 导入后不可编辑", "确认文件解析为图层；普通 SVG、平台 SVG 和内置图元均应支持编辑。"),
        ("SVG 导入后偏大或偏小", "比较 viewBox/use/symbol，并检查端子数量元数据；不要用连接线存在与否判断尺寸。"),
        ("SVG 导出包含端子辅助线", "使用元件右键菜单导出，而不是对整个模型截图或导出。"),
        ("右键菜单偏离鼠标", "确认浏览器缩放为正常值；点击空白关闭后重新右键。"),
        ("切换页面时菜单不消失", "单击其他标签或空白区域；若仍存在，刷新后检查当前版本。"),
        ("关闭窗口没有保存提示", "确认底部存在编辑标识；只选择对象不会产生脏状态。"),
        ("内置图元重启后恢复", "检查用户自定义覆盖是否写入后台以及加载接口是否正常。"),
        ("再次导出没有打开上次目录", "检查 localStorage、目录是否仍存在以及是否使用本地原生导出。"),
        ("浏览器导出直接下载", "本地原生对话框不可用时会回退；浏览器下载目录受浏览器安全策略控制。"),
    ]
    add_table(doc, ["现象", "定位与处理"], troubleshooting, [3000, 6360], caption="常见问题速查")
    add_section(doc, "17.1 推荐的诊断顺序", [
        "先确认对象层级和当前模式，再检查是否已保存，然后检查数据接口和渲染。对于 UI 问题，不要只看按钮文字，应沿“选择对象—形成有效定义—写入草稿—保存后台—重新加载”链路定位。",
    ])
    add_steps(doc, [
        "确认当前选择的是类别库、类、派生类、元件还是模型实例。",
        "确认界面处于类模式、元件模式或模型模式。",
        "检查底部编辑标识和对应保存按钮。",
        "重新选择对象，观察左侧树、右侧表和顶部名称是否一致。",
        "查看后台服务和模型库连接，确认不是接口未加载。",
        "使用最小示例复现，再比较正常对象与异常对象的字段和 SVG 元数据。",
    ])

    add_heading(doc, "附录 A 字段、对象和保存位置速查", 1)
    add_table(doc, ["对象", "主要内容", "保存入口", "继承/引用关系"], [
        ["类别库", "类的组织与能源分类", "类导入导出/后台", "包含基类"],
        ["基类", "公共参数、量测、端子能源属性", "保存类定义", "被派生类继承"],
        ["派生类", "差异参数和差异量测", "保存类定义", "继承基类，不重复字段"],
        ["元件", "图元、状态、端子显示、中英文名", "保存元件定义", "引用所属类语义"],
        ["模型实例", "位置、连接、参数值、量测值", "主界面保存", "引用元件和类定义"],
        ["模板", "多个实例和连接的组合", "模板库保存/导入导出", "用于重复放置"],
    ], [1500, 3000, 2100, 2760], caption="对象层级和保存位置")
    add_table(doc, ["字段", "含义", "默认来源", "是否可编辑"], [
        ["idx", "实例序号", "基类公共字段", "通常由系统维护"],
        ["name", "实例名称", "基类公共字段", "是"],
        ["run_stat", "运行/停运状态", "基类数字枚举", "是"],
        ["dev_type", "所属类英文名", "元件所属类", "是"],
        ["node / i_node / j_node", "端子拓扑节点", "端子定义和连接", "随拓扑维护"],
        ["control_type", "PV/PQ/PH 等控制方式", "设备类定义", "是"],
        ["regable", "是否可调", "数字枚举", "是"],
        ["p_set / q_set / v_set", "控制设定值", "设备类定义", "是"],
    ], [2100, 2600, 2500, 2160], caption="常用字段速查")

    add_heading(doc, "附录 B 操作验收清单", 1)
    add_section(doc, "B.1 类定义验收", ["完成类维护后，按下列清单逐项确认。"], bullets=[
        "基类公共字段存在且英文名唯一。",
        "派生类没有重复定义基类字段。",
        "端子数量、能源属性和 topo 字段一致。",
        "参数取值类型和默认值合理，regable 等枚举类型正确。",
        "量测关联字段属于有效参数集合。",
        "保存后重新选择该类，数据仍然存在。",
    ])
    add_section(doc, "B.2 元件定义验收", ["完成图元维护后，确认图形、端子和持久化。"], bullets=[
        "中文名、英文名和所属类正确。",
        "主体图形位于编辑边界内，没有裁切。",
        "所有图层都可选择、移动、修改和删除。",
        "端子显示位置与所属类端子数量一致。",
        "导出 SVG 不包含端子和端子连接线。",
        "保存并重启 WEB 后，内置或自定义修改仍然存在。",
    ])
    add_section(doc, "B.3 模型交付验收", ["交付前执行模型级检查。"], bullets=[
        "方案、模型、单位和功率基值正确。",
        "模型表字段与类定义有效集合一致。",
        "dev_type、控制类型和设定值符合业务要求。",
        "量测完整且关联字段有效。",
        "拓扑无悬空、冲突和异常节点。",
        "导出目录、文件名、扩展名和文件内容已复核。",
    ])

    add_heading(doc, "附录 C 维护边界与版本管理", 1)
    add_section(doc, "C.1 修改前的备份", [
        "批量调整类字段、端子数量或内置图元前，先导出相关类定义或备份后台数据。元件 SVG 适合备份图形，但不能替代完整类导出，因为 SVG 不包含参数、量测和端子语义。",
    ])
    add_section(doc, "C.2 升级后的回归检查", [
        "程序更新后，抽查典型基类、派生类和元件。推荐至少检查 ACGenerator、ACRealBs、一个多端设备、一个氢能设备和一个热能设备；同时验证类参数、量测、图元持久化、复制粘贴、SVG 导入导出和最近目录。",
    ])
    add_section(doc, "C.3 数据责任边界", [
        "类维护人员负责字段、量测和端子语义；图元维护人员负责图形和端子显示；模型人员负责实例值和拓扑；接口人员负责 E 文件和外部映射。多人协作时，先明确修改对象和保存边界，可以显著减少相互覆盖。",
    ])
    add_note(doc, "版本标识", "本手册适用于 main@1267a934。后续若界面或数据契约发生变化，应更新截图、字段速查和排障表，并重新执行 DOCX 逐页渲染检查。", kind="note")

    add_heading(doc, "附录 D 常用任务入口与保存对象速查", 1)
    add_para(doc, "当界面上同时出现类、元件和模型实例时，可先用下表判断操作入口和实际保存对象。若编辑结果在重新选择后消失，优先检查是否使用了对应层级的保存入口。")
    add_table(doc, ["常用任务", "主要入口", "实际保存对象", "完成后的最小复核"], [
        ["创建基类", "类别库节点右键 → 新建类", "基类定义", "公共字段、端子和 topo 字段自动生成"],
        ["创建派生类", "基类节点右键 → 新建派生类", "派生类差异定义", "继承字段可见，但未在派生类重复保存"],
        ["创建元件", "基类或派生类节点右键 → 新建元件", "元件定义", "左树立即出现记录，右侧保持元件编辑模式"],
        ["复制并跨类粘贴", "元件节点右键复制，目标类节点右键粘贴", "新的元件记录", "名称查重通过，原图层完整且仍可编辑"],
        ["导入 SVG", "元件节点右键 → 导入 SVG 为图元", "当前元件图形", "主体尺寸合理，图层可以增、删、改"],
        ["导出 SVG", "元件节点右键 → 导出图元为 SVG", "外部 SVG 文件", "文件不含端子和端子连接线"],
        ["编辑模型实例", "画布或右侧模型表", "当前模型", "有效字段集合、dev_type、量测和拓扑一致"],
    ], [1900, 2500, 2100, 2860], caption="常用任务、入口和保存边界")

    return doc


def audit_document(doc: Document):
    assert len(doc.sections) == 1
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert all(round(value.inches, 2) == 1.0 for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin))
    assert len(doc.inline_shapes) >= 20
    assert len(doc.tables) >= 10
    assert any(p.style.name == "Heading 1" for p in doc.paragraphs)
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa"
        assert table._tbl.tblPr.find(qn("w:tblInd")) is not None


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    OPTIMIZED_DIR.mkdir(parents=True, exist_ok=True)
    create_diagrams()
    doc = build_manual()
    audit_document(doc)
    doc.save(OUTPUT_PATH)
    print(f"OUTPUT={OUTPUT_PATH}")
    print(f"PARAGRAPHS={len(doc.paragraphs)}")
    print(f"TABLES={len(doc.tables)}")
    print(f"IMAGES={len(doc.inline_shapes)}")


FIGURE_COUNTER = 0
TABLE_COUNTER = 0
BULLET_ABSTRACT_ID = 900
DECIMAL_ABSTRACT_ID = 901


if __name__ == "__main__":
    main()
